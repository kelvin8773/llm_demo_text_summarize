# utils/ingest.py - Enhanced document loading with error handling
import docx
from PyPDF2 import PdfReader
import logging

logger = logging.getLogger(__name__)


def load_document(file):
    """
    Load text content from various document formats.
    
    Args:
        file: File object (uploaded file or file-like object)
        
    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: For unsupported file types or empty documents
        Exception: For file processing errors
    """
    if not file:
        raise ValueError("No file provided")
    
    file_name = getattr(file, 'name', 'unknown')
    file_size = getattr(file, 'size', 0)
    
    # Check file size (limit to 10MB)
    if file_size > 10 * 1024 * 1024:
        raise ValueError(f"File too large ({file_size / (1024*1024):.1f}MB). Maximum size is 10MB.")
    
    try:
        if file_name.lower().endswith(".pdf"):
            return _load_pdf(file)
        elif file_name.lower().endswith(".txt"):
            return _load_txt(file)
        elif file_name.lower().endswith(".docx"):
            return _load_docx(file)
        else:
            raise ValueError(f"Unsupported file type: {file_name}. Supported formats: PDF, TXT, DOCX")
            
    except Exception as e:
        logger.error(f"Error loading document {file_name}: {str(e)}")
        raise Exception(f"Failed to load document: {str(e)}")


def _load_pdf(file):
    """Load text from PDF file with enhanced error handling."""
    try:
        reader = PdfReader(file)
        
        if len(reader.pages) == 0:
            raise ValueError("PDF file appears to be empty or corrupted")
        
        text_parts = []
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
                else:
                    logger.warning(f"Page {i+1} appears to be empty or contains only images")
            except Exception as e:
                logger.warning(f"Could not extract text from page {i+1}: {str(e)}")
                continue
        
        if not text_parts:
            raise ValueError("No readable text found in PDF. The document may contain only images or be password-protected.")
        
        text = " ".join(text_parts)
        
        if len(text.strip()) < 50:
            raise ValueError("PDF contains very little text. Please ensure the document has readable text content.")
        
        return text
        
    except Exception as e:
        if "password" in str(e).lower():
            raise ValueError("PDF appears to be password-protected. Please provide an unlocked version.")
        elif "corrupted" in str(e).lower():
            raise ValueError("PDF file appears to be corrupted. Please try a different file.")
        else:
            raise Exception(f"Error reading PDF: {str(e)}")


def _load_txt(file):
    """Load text from TXT file with encoding detection."""
    try:
        # Try UTF-8 first
        try:
            file.seek(0)  # Reset file pointer
            text = file.read().decode("utf-8")
        except UnicodeDecodeError:
            # Try other common encodings
            file.seek(0)
            try:
                text = file.read().decode("latin-1")
            except UnicodeDecodeError:
                file.seek(0)
                text = file.read().decode("cp1252")
        
        if not text or len(text.strip()) < 10:
            raise ValueError("Text file appears to be empty or contains very little content")
        
        return text.strip()
        
    except Exception as e:
        raise Exception(f"Error reading text file: {str(e)}")


def _load_docx(file):
    """Load text from DOCX file with error handling."""
    try:
        doc = docx.Document(file)
        
        if not doc.paragraphs:
            raise ValueError("DOCX file appears to be empty")
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        if not text_parts:
            raise ValueError("No readable text found in DOCX file")
        
        text = " ".join(text_parts)
        
        if len(text.strip()) < 50:
            raise ValueError("DOCX contains very little text. Please ensure the document has readable content.")
        
        return text
        
    except Exception as e:
        if "corrupted" in str(e).lower():
            raise ValueError("DOCX file appears to be corrupted. Please try a different file.")
        else:
            raise Exception(f"Error reading DOCX file: {str(e)}")
